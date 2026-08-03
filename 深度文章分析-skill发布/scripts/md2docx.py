#!/usr/bin/env python3
"""
Markdown → DOCX converter.
Produces a Word/WPS-compatible .docx with proper styling:
  - H1/H2/H3 with theme colors
  - Bold, lists, blockquotes, code blocks, tables
  - Page margins, fonts, footer
"""

import sys
import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ---- Style constants ----
COLOR_H1 = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_H2 = RGBColor(0x25, 0x63, 0xEB)  # blue accent
COLOR_H3 = RGBColor(0x33, 0x33, 0x33)
COLOR_BODY = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_QUOTE = RGBColor(0x55, 0x55, 0x55)
COLOR_CODE_BG = RGBColor(0xF5, 0xF5, 0xF5)
COLOR_CODE_TEXT = RGBColor(0x33, 0x33, 0x33)
COLOR_TABLE_HEADER_BG = "2563EB"
COLOR_TABLE_HEADER_FG = "FFFFFF"
COLOR_TABLE_ALT = "F0F6FF"
COLOR_FOOTER = RGBColor(0x99, 0x99, 0x99)
COLOR_LINE = RGBColor(0xDD, 0xDD, 0xDD)
COLOR_QUOTE_BAR = RGBColor(0x25, 0x63, 0xEB)

FONT_NAME = "Microsoft YaHei"
FONT_NAME_MONO = "Consolas"
FONT_SIZE_BODY = 10.5
FONT_SIZE_H1 = 18
FONT_SIZE_H2 = 14
FONT_SIZE_H3 = 12
FONT_SIZE_CODE = 9
FONT_SIZE_FOOTER = 8.5


def setup_styles(doc):
    """Configure document-wide defaults via XML."""
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE_BODY)
    style.font.color.rgb = COLOR_BODY
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(4)

    # Set font for East Asian text too
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rpr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)


def parse_inline_markdown(text):
    """
    Parse inline markdown (**bold**, `code`, links) into segments.
    Returns list of (text, is_bold, is_code, is_italic).
    """
    segments = []
    # Combined pattern: **bold** | `code` | *italic* | [text](url) 
    pattern = re.compile(r'\*\*(.+?)\*\*|`(.+?)`|\*(.+?)\*|\[(.+?)\]\(.+?\)')
    last_end = 0
    for m in pattern.finditer(text):
        if m.start() > last_end:
            segments.append((text[last_end:m.start()], False, False, False))
        if m.group(1):  # **bold**
            segments.append((m.group(1), True, False, False))
        elif m.group(2):  # `code`
            segments.append((m.group(2), False, True, False))
        elif m.group(3):  # *italic*
            segments.append((m.group(3), False, False, True))
        elif m.group(4):  # [text](url) — just use the text
            segments.append((m.group(4), False, False, False))
        last_end = m.end()
    if last_end < len(text):
        segments.append((text[last_end:], False, False, False))
    return segments


def add_run(paragraph, text, bold=False, italic=False, font_name=None, 
            font_size=None, color=None):
    """Add a run with specified formatting."""
    run = paragraph.add_run(text)
    run.font.name = font_name or FONT_NAME
    run.font.size = Pt(font_size or FONT_SIZE_BODY)
    if color:
        run.font.color.rgb = color
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    # Set East Asian font
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rpr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), FONT_NAME)
    return run


def add_rich_paragraph(doc, text):
    """Add a paragraph with inline formatting (bold, code, italic)."""
    segments = parse_inline_markdown(text)
    if not segments:
        return doc.add_paragraph()

    p = doc.add_paragraph()
    for seg_text, is_bold, is_code, is_italic in segments:
        if is_code:
            add_run(p, seg_text, font_name=FONT_NAME_MONO, 
                    font_size=FONT_SIZE_CODE, color=COLOR_CODE_TEXT)
        else:
            add_run(p, seg_text, bold=is_bold, italic=is_italic)
    return p


def add_heading(doc, text, level):
    """Add a heading with custom styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.bold = True

    if level == 1:
        run.font.size = Pt(FONT_SIZE_H1)
        run.font.color.rgb = COLOR_H1
        p.paragraph_format.space_before = Pt(28)
        p.paragraph_format.space_after = Pt(8)
        # Bottom border
        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="2563EB"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(FONT_SIZE_H2)
        run.font.color.rgb = COLOR_H2
        p.paragraph_format.space_before = Pt(22)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(FONT_SIZE_H3)
        run.font.color.rgb = COLOR_H3
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)

    # East Asian font
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rpr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), FONT_NAME)
    return p


def add_bullet_list(doc, items):
    """Add bullet list items."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        add_run(p, item)


def add_numbered_list(doc, items):
    """Add numbered list items."""
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.clear()
        add_run(p, item)


def add_blockquote(doc, lines):
    """Add a blockquote with left bar (indent + accent border)."""
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_QUOTE
        run.font.italic = True
        # Left indent
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        # Left border (accent)
        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:left w:val="single" w:sz="12" w:space="6" w:color="2563EB"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
        # Shading
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F0F6FF"/>')
        pPr.append(shd)


def add_code_block(doc, lines):
    """Add a code block with monospace font and grey background."""
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = FONT_NAME_MONO
        run.font.size = Pt(FONT_SIZE_CODE)
        run.font.color.rgb = COLOR_CODE_TEXT
        # Left indent
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        # Grey shading
        pPr = p._element.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F0F0F0"/>')
        pPr.append(shd)


def add_horizontal_rule(doc):
    """Add a horizontal rule (thin bottom border on empty paragraph)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="DDDDDD"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_table(doc, md_table_text):
    """Parse a simple markdown table and add it to docx."""
    lines = md_table_text.strip().split('\n')
    if len(lines) < 2:
        return

    # Parse header
    header_cells = [c.strip() for c in lines[0].split('|') if c.strip()]
    if len(lines) >= 3:
        data_rows = []
        for line in lines[2:]:
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                data_rows.append(cells)
    else:
        data_rows = []

    if not header_cells:
        return

    ncols = len(header_cells)
    table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
    table.style = 'Table Grid'
    table.autofit = True

    # Header row
    for i, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
        # Blue background
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        shd = parse_xml(
            f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="2563EB"/>'
        )
        tcPr.append(shd)

    # Data rows
    for ri, row_data in enumerate(data_rows):
        for ci, cell_text in enumerate(row_data):
            if ci >= ncols:
                break
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(9.5)
            # Alternate row color
            if ri % 2 == 1:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                shd = parse_xml(
                    f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F0F6FF"/>'
                )
                tcPr.append(shd)

    doc.add_paragraph()  # spacing after table


def md_to_docx(md_path, docx_path):
    """Convert markdown file to styled .docx."""
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = Document()
    setup_styles(doc)

    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []
    in_blockquote = False
    quote_lines = []

    while i < len(lines):
        line = lines[i]

        # Code block (fenced) — handle as raw text block
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                if code_block_lines:
                    add_code_block(doc, code_block_lines)
                in_code_block = False
                code_block_lines = []
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Table detection
        if '|' in line and line.strip().startswith('|'):
            table_lines.append(line)
            in_table = True
            i += 1
            continue
        elif in_table:
            # End of table
            if table_lines:
                add_table(doc, '\n'.join(table_lines))
            in_table = False
            table_lines = []
            continue  # re-process this line

        # Blockquote
        if line.startswith('> '):
            quote_lines.append(line[2:])
            i += 1
            continue
        elif quote_lines:
            add_blockquote(doc, quote_lines)
            quote_lines = []
            continue

        # Horizontal rule
        if re.match(r'^---+$', line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,3})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            add_heading(doc, text, level)
            i += 1
            continue

        # Bullet list
        if re.match(r'^[\s]*[-*+]\s+', line):
            items = []
            while i < len(lines) and re.match(r'^[\s]*[-*+]\s+', lines[i]):
                item_text = re.sub(r'^[\s]*[-*+]\s+', '', lines[i])
                items.append(item_text)
                i += 1
            add_bullet_list(doc, items)
            continue

        # Numbered list
        if re.match(r'^\s*\d+[\.\)]\s+', line):
            items = []
            while i < len(lines) and re.match(r'^\s*\d+[\.\)]\s+', lines[i]):
                item_text = re.sub(r'^\s*\d+[\.\)]\s+', '', lines[i])
                items.append(item_text)
                i += 1
            add_numbered_list(doc, items)
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Plain paragraph — merge consecutive lines
        para = line.strip()
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
                r'^(#{1,3}\s|[\s]*[-*+]\s|\s*\d+[\.\)]\s|> |---+$)', lines[i]):
            para += ' ' + lines[i].strip()
            i += 1
        add_rich_paragraph(doc, para)

    # Flush remaining blocks
    if quote_lines:
        add_blockquote(doc, quote_lines)
    if code_block_lines and in_code_block:
        add_code_block(doc, code_block_lines)
    if table_lines:
        add_table(doc, '\n'.join(table_lines))

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run('—— 由 WorkBuddy 深度阅读分析大师生成 ——')
    run.font.size = Pt(FONT_SIZE_FOOTER)
    run.font.color.rgb = COLOR_FOOTER

    doc.save(docx_path)
    size = os.path.getsize(docx_path)
    print(f"DOCX generated: {docx_path} ({size} bytes)")
    return docx_path


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python md2docx.py <input.md> <output.docx>")
        sys.exit(1)
    md_to_docx(sys.argv[1], sys.argv[2])
